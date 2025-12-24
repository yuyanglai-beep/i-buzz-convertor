import gradio as gr
from docx import Document
from html import escape
import tempfile
import re
import os

# 🖼️ Image tool deps
from PIL import Image
import io

# ==========================================
# 1. 核心功能 (Core Logic) - Footer 區
# ==========================================

FOOTER_VOC = """<div class="cta-card-wrapper"><div class="cta-card footer-card"><p class="footer-text"><em><strong class="footer-strong">i-Buzz 行業資料庫</strong><span class="footer-normal"> 提供各產業完整的品牌、產品資訊與熱門話題內容。您只需告訴我們想了解的產業領域，無需設定繁複關鍵字，即可快速取得可靠的行業口碑洞察。如需更深入的數據應用與分析服務，歡迎填寫表單與我們聯繫。</span></em></p><p class="footer-btn-area"><a href="https://www.i-buzz.com.tw/user/contact/" target="_blank"><img src="/uploads/industry_img/1712050466.png" alt="CTA" width="200" height="65"></a></p></div></div><style>.cta-card-wrapper { display: flex; justify-content: center; margin: 40px 0; } .footer-card { max-width: 720px; background: linear-gradient(180deg, #f4e9dc 0%, #fffdf9 90%); padding: 38px 45px; border-radius: 18px; border: none; box-shadow: 0 10px 28px rgba(140, 110, 70, 0.18); } .footer-text { font-size: 16px; line-height: 1.85; margin-bottom: 28px; color: #5a371e; } .footer-normal { color: #5a371e; } .footer-strong { color: #7b4a21; font-weight: 700; } .footer-btn-area { text-align: center; }</style>"""
FOOTER_TREND = """<div class="cta-card-wrapper"><div class="cta-card footer-card"><p class="footer-text"><em><span class="footer-normal">i-Buzz為台灣首間網路口碑研究中心，累積超過16年的跨產業口碑分析經驗， 提供客戶 </span> <strong class="footer-strong">i-Buzz VOC+ 產業口碑數據庫、商業策略分析</strong> <span class="footer-normal"> 、消費者輪廓洞察及口碑行銷優化等全方位服務。 歡迎填寫表單，讓專業的團隊為您服務！ </span> </em></p><p class="footer-btn-area"><a href="https://www.i-buzz.com.tw/user/contact/" target="_blank"><img alt="CTA" height="65" src="/uploads/industry_img/1712050466.png" width="200" /> </a></p></div></div><style type="text/css">.cta-card-wrapper { display: flex; justify-content: center; margin: 40px 0; } .footer-card { max-width: 720px; background: linear-gradient(180deg, #f4e9dc 0%, #fffdf9 90%); padding: 38px 45px; border-radius: 18px; border: none; box-shadow: 0 10px 28px rgba(140, 110, 70, 0.18); } .footer-text { font-size: 16px; color: #5a371e; line-height: 1.85; margin-bottom: 28px; } .footer-normal { color: #5a371e; } .footer-strong { color: #7b4a21; font-weight: 700; } .footer-btn-area { text-align: center; }</style>"""
FOOTER_AK = """
<style>
  .cta-section {
    text-align: center;
    color: #2c7a7b;
    font-family: "Noto Sans TC", sans-serif;
    margin: 20px auto 40px;
    line-height: 1.6;
    max-width: 780px;
  }

  .cta-block {
    margin-bottom: 32px; /* 兩段 CTA 之間的距離更短 */
  }

  .cta-star {
    font-size: 18px;
    margin-right: 4px;
  }

  .cta-title {
    font-size: 18px;
    font-weight: 700;
    margin-bottom: 10px; /* 標題與文字距離縮短 */
  }

  .cta-desc {
    font-size: 15px;
    margin-bottom: 18px; /* 變緊湊 */
  }

  .cta-btn {
    display: inline-block;
    padding: 10px 22px;
    background: #2c7a7b;
    color: white;
    border-radius: 6px;
    text-decoration: none;
    font-size: 15px;
    font-weight: 600;
    transition: 0.25s;
  }

  .cta-btn:hover {
    background: #225f61;
  }
</style>
<style type="text/css">.cta-section {
    text-align: center;
    color: #2c7a7b; 
    font-family: "Noto Sans TC", sans-serif;
    margin: 40px auto;
    line-height: 1.8;
    max-width: 820px;
  }

  .cta-star {
    font-size: 22px;
    margin-right: 6px;
  }

  .cta-text {
    font-size: 18px;
    font-weight: 600;
    margin: 20px 0 10px;
  }

  .cta-desc {
    font-size: 16px;
    margin-bottom: 28px;
  }

  .cta-btn {
    display: inline-block;
    padding: 10px 28px;
    margin: 8px 0 25px;
    background: #2c7a7b;
    color: white;
    border-radius: 6px;
    text-decoration: none;
    font-size: 16px;
    font-weight: 600;
    transition: 0.25s;
  }

  .cta-btn:hover {
    background: #225f61;
  }
</style>
<style type="text/css">.cta-one {
    text-align: center;
    color: #2c7a7b;
    font-family: "Noto Sans TC", sans-serif;
    max-width: 800px;
    margin: 40px auto;
    line-height: 1.7;
  }

  .cta-one strong {
    font-size: 18px;
  }

  .cta-btn-wrap {
    margin-top: 16px;
    display: flex;
    justify-content: center;
    gap: 14px;
    flex-wrap: wrap;
  }

  .cta-btn2 {
    display: inline-block;
    padding: 10px 22px;
    background: #2c7a7b;
    color: white;
    border-radius: 6px;
    text-decoration: none;
    font-size: 15px;
    font-weight: 600;
    transition: 0.25s;
  }

  .cta-btn2:hover {
    background: #225f61;
  }
</style>
<style type="text/css">.cta-box {
    border: 1.5px solid #c7d8d8; /* 外框顏色 */
    border-radius: 10px;
    padding: 28px 22px;
    max-width: 820px;
    margin: 40px auto;
    background: #f9fcfc; /* 淡淡底色，讓 CTA 更突出 */
  }

  .cta-one {
    text-align: center;
    color: #2c7a7b;
    font-family: "Noto Sans TC", sans-serif;
    line-height: 1.7;
    font-size: 16px;
  }

  .cta-one strong {
    font-size: 18px;
  }

  .cta-btn-wrap {
    margin-top: 18px;
    display: flex;
    justify-content: center;
    gap: 14px;
    flex-wrap: wrap;
  }

  .cta-btn2 {
    display: inline-block;
    padding: 10px 22px;
    background: #2c7a7b;
    color: white;
    border-radius: 6px;
    text-decoration: none;
    font-size: 15px;
    font-weight: 600;
    transition: 0.25s;
  }

  .cta-btn2:hover {
    background: #225f61;
  }
</style>
</strong></p>

<div class="cta-box">
<div class="cta-one">
<p><strong>⭐ <strong>「AsiaKOL 網紅專案式顧問服務」</strong></strong><strong style="font-size: 18px; color: rgb(44, 122, 123); font-family: &quot;Noto Sans TC&quot;, sans-serif; text-align: center; background-color: rgb(249, 252, 252);"><strong>，</strong></strong><strong><strong>從網紅精準篩選、創意內容企劃到專案執行與監測，全程由專業團隊一手打造。</strong><br />
若您想了解更多服務內容，或希望由專人協助規劃合作，歡迎點擊下方： </strong></p>

<div class="cta-btn-wrap"><a class="cta-btn2" href="https://www.asiakol.com/page/view/service/project" target="_blank">服務介紹</a> <a class="cta-btn2" href="https://www.asiakol.com/page/view/contact-us" target="_blank">填寫需求單</a></div>
</div>
</div>

"""
FOOTER_FF = """
<hr />
<p style="margin: 0px; padding: 0px;">
  <span id="docs-internal-guid-008b77ae-7fff-c582-a34d-ac1dfa7fefd7">
    <span style="font-weight: 700; font-family: Arial, sans-serif; font-size: 12pt; color: rgb(255, 255, 255); background-color: rgb(0, 0, 128);">FANS FEED 品牌頻道經營</span>
  </span>
</p>
<p style="margin: 14pt 0px; line-height: 1.2;">
  <span id="docs-internal-guid-008b77ae-7fff-c582-a34d-ac1dfa7fefd7">
    <span style="font-family: REM, sans-serif; font-size: 12pt; color: rgb(85, 85, 85);"> ⭐ </span>
    <span style="font-family: Arial, sans-serif; font-size: 12pt; color: rgb(85, 85, 85);"> </span>
    <span style="font-family: Arial, sans-serif; font-size: 12pt; color: rgb(0, 0, 128);">品牌小編努力發文，成效卻不見起色嗎？你需要經驗豐富的專業小編團隊，為你管理官方社群頻道，以數據分析及深度觀察達到內容精采度與宣傳成效&nbsp;►&nbsp;</span>
    <a href="https://fansfeed.com.tw/cultivateserve_p1_1" style="text-decoration-line: none; color: rgb(66, 174, 251);" target="_blank">
      <span style="font-weight: 700; font-family: Arial, sans-serif; font-size: 12pt; color: rgb(255, 255, 255); background-color: rgb(0, 128, 128); text-decoration-line: underline;">了解更多</span>
    </a>
  </span>
</p>
<p style="margin: 14pt 0px 0pt; line-height: 1.2;">
  <span style="font-size: 12pt; font-family: REM, sans-serif; color: rgb(85, 85, 85);"> ⭐ </span>
  <span style="font-size: 12pt; font-family: Arial, sans-serif; color: rgb(85, 85, 85);"> </span>
  <span style="font-size: 12pt; font-family: Arial, sans-serif; color: rgb(0, 0, 128);">填寫需求單，將有專業團隊為你服務&nbsp;►</span>
  <span style="color: rgb(255, 255, 255);"><span style="font-family: Arial, sans-serif;">&nbsp;</span></span>
  <span style="font-size: 12pt; font-family: Arial, sans-serif; font-weight: 700; background-color: rgb(0, 128, 128);">
    <span>
      <span>
        <a href="https://fansfeed.com.tw/index#CBArrow" style="text-decoration-line: none; color: rgb(66, 174, 251);" target="_blank">
          <span style="color: rgb(255, 255, 255);">立即填寫</span>
        </a>
      </span>
    </span>
  </span>
</p>
"""
FOOTER_THREADS = """<div class="cta-card-wrapper"><div class="cta-card"><p class="cta-title"><strong>Threads 爆發力強、紅利正旺！</strong></p><p class="cta-subtitle">在高流量、高競爭的環境裡，品牌只有一次被看見的機會。</p><p class="cta-desc">i-Buzz Threads 行銷服務，讓你的內容更有話題、更容易衝上熱度高點。</p><p class="cta-highlight"><span class="highlight-light">想讓品牌成為下一個爆紅案例？</span><span class="highlight-bold">和我們聊聊吧。</span></p><p class="cta-btn-area"><a href="https://www.i-buzz.com.tw/article/threadsmarketing#treads_sec_4" target="_blank"><img src="https://www.i-buzz.com.tw/uploads/industry_img/1712050466.png" alt="CTA" width="220" height="70"></a></p></div></div><style type="text/css">.cta-card-wrapper { display: flex; justify-content: center; margin: 40px 0; } .cta-card { max-width: 720px; background: linear-gradient(180deg, #f6f3ff 0%, #ffffff 85%); padding: 40px 45px; border-radius: 22px; box-shadow: 0 14px 36px rgba(80, 60, 140, 0.15); border: 1px solid #ece8ff; } .cta-title { font-size: 30px; font-weight: 800; color: #4f17b1; margin: 0 0 4px; line-height: 1.3; } .cta-subtitle { font-size: 18px; color: #4f17b1; margin: 0 0 28px; line-height: 1.45; } .cta-desc { font-size: 17px; color: #7c6af2; line-height: 1.75; margin-bottom: 22px; } .cta-highlight { font-size: 19px; line-height: 1.7; margin-bottom: 32px; } .highlight-light { color: #8e7dfa; } .highlight-bold { color: #4f17b1; font-weight: 700; } .cta-btn-area { text-align: center; margin-top: 10px; }</style>"""

CATEGORY_TO_FOOTER_HTML = {
    "🔵 數據分析解方": FOOTER_VOC,
    "🔷 產業口碑數據": FOOTER_TREND,
    "🟦 消費者洞察": FOOTER_VOC,
    "🩷 網紅行銷策略": FOOTER_AK,
    "🟡 社群粉絲團健檢": FOOTER_FF,
    "🟣 Threads 行銷服務": FOOTER_THREADS
}

# ==========================================
# 內容控制項移除（Word 表單欄位 → 一般文字）
# ==========================================

def remove_content_controls(doc):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    sdt_nodes = list(doc.element.body.xpath('.//*[local-name()="sdt"]'))
    for sdt in sdt_nodes:
        parent = sdt.getparent()
        if parent is None:
            continue
        sdt_content = sdt.xpath('./*[local-name()="sdtContent"]')
        if sdt_content:
            sdt_content = sdt_content[0]
            insert_at = parent.index(sdt)
            for child in list(sdt_content):
                parent.insert(insert_at, child)
                insert_at += 1
            parent.remove(sdt)
        else:
            texts = sdt.xpath('.//*[local-name()="t"]')
            combined = ''.join(t.text or '' for t in texts)
            if combined.strip():
                run = OxmlElement(qn('w:r'))
                t = OxmlElement(qn('w:t'))
                t.text = combined
                run.append(t)
                parent.insert(parent.index(sdt), run)
            parent.remove(sdt)

# ==========================================
# 段落 → HTML（保留超連結）
# ==========================================

def paragraph_to_html_with_links(para):
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    runs_html = []
    for child in para._element:
        tag = child.tag.split('}')[-1]
        if tag == 'hyperlink':
            rel_id = child.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
            href = para.part.rels[rel_id].target_ref if rel_id and rel_id in para.part.rels else None
            link_text = ''.join(t.text or '' for t in child.findall('.//w:t', ns))
            if href:
                runs_html.append(f'<a href="{escape(href)}" target="_blank">{escape(link_text)}</a>')
            else:
                runs_html.append(escape(link_text))
        else:
            text = ''.join(t.text or '' for t in child.findall('.//w:t', ns))
            if text:
                runs_html.append(escape(text))
    if not runs_html:
        runs_html.append(escape(para.text or ''))
    return ''.join(runs_html)

# ==========================================
# 表格 → HTML
# ==========================================

def table_to_html(table):
    html = ['<table style="border-collapse:collapse; table-layout:auto; border:1px solid #ccc; margin-left:0; margin-right:auto;">']
    for r_index, row in enumerate(table.rows):
        html.append('<tr>')
        for cell in row.cells:
            cell_content = []
            for para in cell.paragraphs:
                cell_content.append(paragraph_to_html_with_links(para))
            cell_html = "<br>".join(cell_content)
            if r_index == 0:
                cell_html = f"<strong>{cell_html}</strong>"
            html.append(f'<td style="border:1px solid #ccc; padding:6px; vertical-align:top;">{cell_html}</td>')
        html.append('</tr>')
    html.append('</table>')
    return ''.join(html)

# ==========================================
# 判斷是否為編號列點 (新增)
# ==========================================
def is_numbered_list(para):
    """檢查是否為 Word 內建編號或手打數字開頭"""
    # 1. 檢查 Word 內建編號屬性
    pPr = para._element.get_or_add_pPr()
    if pPr.xpath('./w:numPr'):
        return True
    # 2. 檢查手打數字 (例如: 1. 內容 或 1、內容)
    text = para.text.strip()
    if re.match(r"^\d+[\.\s、．]+", text):
        return True
    return False

# ==========================================================
# 📌 URL → iframe 轉換功能
# ==========================================================

def is_pure_url(text: str) -> bool:
    """判斷段落是否為『單獨只有 URL』"""
    if not text:
        return False
    text = text.strip()
    # 必須整段就是一個 http(s):// 開頭的網址
    return bool(re.fullmatch(r"https?://\S+", text))


def extract_embed_url(text: str):
    """從一段文字中抓出支援平台的網址（IG / Threads / FB / YouTube）"""
    patterns = [
        r"https?://(?:www\.)?instagram\.com/[^\s]+",
        r"https?://(?:www\.)?threads\.net/[^\s]+",
        r"https?://(?:www\.)?facebook\.com/[^\s]+",
        r"https?://(?:www\.)?youtu\.be/[^\s]+",
        r"https?://(?:www\.)?youtube\.com/[^\s]+",
    ]
    for p in patterns:
        m = re.search(p, text)
        if m:
            return m.group(0)
    return None


def convert_url_to_iframe(url: str):
    """中度智慧版：依平台 + 類型自動選高度"""

    # ---------------------------------------------------
    # Instagram 判斷（/p/ = 圖文、/reel/ = 短影音、/tv/ = IGTV）
    # ---------------------------------------------------
    if "instagram.com" in url:
        clean = url.split("?")[0].rstrip("/")

        # 類型判斷
        if "/reel/" in clean:
            height = 800   # Reels
        elif "/tv/" in clean:
            height = 800   # IGTV
        else:
            height = 770   # 一般貼文（單圖 / 輪播）

        embed_url = clean + "/embed"

        return f'''
<p>
  <iframe
      src="{embed_url}"
      scrolling="no"
      style="
          width:100%;
          max-width:480px;
          height:{height}px;
          border:0;
          border-radius:14px;
          display:block;
          margin:0;
      ">
  </iframe>
</p>
'''

    # ---------------------------------------------------
    # Threads 判斷（文字 / 圖片 / 影片）
    # ---------------------------------------------------
    if "threads.net" in url or "threads.com" in url:

        # threads.com → threads.net
        url = url.replace("threads.com", "threads.net")

        clean = url.split("?")[0].rstrip("/")
        embed_url = clean + "/embed"

        # 初步判斷（依 URL & embed pattern）
        lower = url.lower()
        if "photo" in lower or "image" in lower:
            height = 580  # 圖片貼文
        elif "video" in lower or "reel" in lower:
            height = 650  # 影片貼文
        else:
            height = 480  # 文字貼文

        return f'''
<p>
  <iframe
      src="{embed_url}"
      scrolling="no"
      style="
          width:100%;
          max-width:480px;
          height:{height}px;
          border:0;
          border-radius:14px;
          display:block;
          margin:0;
      ">
  </iframe>
</p>
'''

    # ---------------------------------------------------
    # YouTube（固定 16:9）
    # ---------------------------------------------------
    if "youtube.com" in url or "youtu.be" in url:
        if "youtu.be" in url:
            vid = url.split("/")[-1]
        else:
            m = re.search(r"v=([^&]+)", url)
            vid = m.group(1) if m else ""

        return f'''
<p>
  <iframe
      src="https://www.youtube.com/embed/{vid}"
      style="
          width:100%;
          max-width:480px;
          height:270px;
          border:0;
          border-radius:14px;
          display:block;
          margin:0;
      "
      allowfullscreen>
  </iframe>
</p>
'''

    # ---------------------------------------------------
    # Facebook（判斷：影片 or 貼文）
    # ---------------------------------------------------
    if "facebook.com" in url:

        lower = url.lower()

        # 有 videos / watch → 影片貼文
        if "videos" in lower or "video" in lower or "watch" in lower:
            height = 900
        else:
            height = 600

        return f'''
<p>
  <iframe
      src="https://www.facebook.com/plugins/post.php?href={url}"
      scrolling="no"
      style="
          width:100%;
          max-width:480px;
          height:{height}px;
          border:0;
          border-radius:14px;
          display:block;
          margin:0;
      ">
  </iframe>
</p>
'''

    # ---------------------------------------------------
    # 不支援的平台
    # ---------------------------------------------------
    return None


# ===================================================
# ✅ TOC：只插在「第一個 H2 前」，加平滑捲動
# ===================================================

def apply_auto_toc_and_smooth(html_list):
    updated = []
    toc = []
    first_h2_index = None

    smooth_css = """<style>html { scroll-behavior: smooth; } .html-container h2, .html-container h3 { scroll-margin-top: 130px; }</style>"""
    updated.append(smooth_css)

    for block in html_list:
        b = (block or "").strip()
        plain_text = re.sub(r"<[^>]*>", "", b) if b else ""

        if b.startswith("<h1"):
            anchor = f"toc-h1-{len(toc)}"
            toc.append((1, plain_text, anchor))
            block = block.replace("<h1>", f"<h1 id='{anchor}'>", 1)

        elif b.startswith("<h2"):
            anchor = f"toc-h2-{len(toc)}"
            toc.append((2, plain_text, anchor))
            block = block.replace("<h2", f"<h2 id='{anchor}'", 1)
            if first_h2_index is None:
                first_h2_index = len(updated)

        elif b.startswith("<h3"):
            anchor = f"toc-h3-{len(toc)}"
            toc.append((3, plain_text, anchor))
            block = block.replace("<h3", f"<h3 id='{anchor}'", 1)

        updated.append(block)

    if first_h2_index is None or not toc:
        return updated

    def is_blank_para(x: str):
        s = (x or "").strip()
        return s in ("<p>&nbsp;</p>", "<p>&nbsp;</p><p>&nbsp;</p>")

    while first_h2_index - 1 >= 0 and is_blank_para(updated[first_h2_index - 1]):
        updated.pop(first_h2_index - 1)
        first_h2_index -= 1

    toc_html = [
        "<div style='margin-top:28px; margin-bottom:12px; padding:12px 0 12px 16px; border-left:4px solid #4f8ef7;'>",
        "<div style='font-size:20px; font-weight:700; margin-bottom:10px; color:#000000;'>文章目錄</div>",
        "<ul style='list-style:none; margin-left: 28px; padding-left: 0; line-height:1.8; font-size:17px; color:#4f8ef7;'>"
    ]

    for level, text, anchor in toc:
        safe_text = escape(text) if text else ""
        if level == 2:
            bullet = "•"
            indent_px = 0
        elif level == 3:
            bullet = "◦"
            indent_px = 18
        else:
            bullet = "•"
            indent_px = 0

        toc_html.append(
            f"<li style='margin:6px 0; padding-left:{indent_px}px; text-indent:-12px;'>"
            f"<span style='display:inline-block; width:12px; opacity:0.7;'>{bullet}</span>"
            f"<a href='#{anchor}' style='color:#4f8ef7; text-decoration:none;'>{safe_text}</a>"
            f"</li>"
        )

    toc_html.append("</ul></div>")

    updated = updated[:first_h2_index] + toc_html + updated[first_h2_index:]
    insert_after_toc = first_h2_index + len(toc_html)
    updated = updated[:insert_after_toc] + ["<p>&nbsp;</p><p>&nbsp;</p>"] + updated[insert_after_toc:]

    return updated

# ==========================================================
# DOCX → HTML 主流程（含 URL → iframe）
# ==========================================================

def docx_to_html_with_links(input_file, category_choice):
    doc = Document(input_file.name)
    remove_content_controls(doc)
    html_output = []
    h1_text = None
    last_was_blank = False
    in_list = False
    elements = list(doc.element.body)
    total = len(elements)

    for idx, element in enumerate(elements):
        tag = element.tag.split('}')[-1]
        next_tag = elements[idx + 1].tag.split('}')[-1] if idx < total - 1 else None
        next_style = ""
        if next_tag == "p":
            for np in doc.paragraphs:
                if np._element == elements[idx + 1]:
                    next_style = np.style.name.lower() if np.style and np.style.name else ''
                    break

        if tag == "p":
            para = next((p for p in doc.paragraphs if p._element == element), None)
            if not para: continue

            text = (para.text or "").strip()
            style = para.style.name.lower() if para.style and para.style.name else ''

            # --- 🚀 [新增邏輯] 處理標號列點 ---
            if is_numbered_list(para) and text:
                if not in_list:
                    html_output.append("<ol>") # 開啟清單
                    in_list = True
                
                content_html = paragraph_to_html_with_links(para)
                # 清除文字中重複的開頭數字 (如 "1. " 或 "2、")，交給 HTML 自動編號
                content_html = re.sub(r"^\d+[\.\s、．]+", "", content_html)
                html_output.append(f"  <li>{content_html}</li>")
                last_was_blank = False
                continue # 跳過後續的一般段落處理
            else:
                if in_list:
                    html_output.append("</ol>") # 遇到非列點文字，關閉清單
                    in_list = False
            # --- [新增邏輯結束] ---

            if not text:
                continue

            # 👉 先處理「整段只有 URL」→ 嘗試轉成 iframe
            if is_pure_url(text):
                url = extract_embed_url(text) or text
                iframe_block = convert_url_to_iframe(url)
                if iframe_block:
                    # 確保 iframe 上方只有 1 個空行
                    if not last_was_blank:
                        html_output.append("<p>&nbsp;</p>")
                    
                    html_output.append(iframe_block)
                    
                    # iframe 下方固定 1 行空行
                    html_output.append("<p>&nbsp;</p>")
                    
                    last_was_blank = True

                    continue  # 這段已處理完，不再用一般段落邏輯

            # 一般段落轉換
            content_html = paragraph_to_html_with_links(para)

            # 移除 Word 自帶 TOC 的編號
            if style.startswith("toc"):
                content_html = re.sub(r"^\s*[\d\.\-\(\)、．]+\s*", "", content_html)

            # H1 當作主標題，只抓文字不輸出
            if 'heading 1' in style:
                if not h1_text:
                    h1_text = text
                continue

            # H2
            if 'heading 2' in style:
                html_output.append('<p>&nbsp;</p>' if last_was_blank else '<p>&nbsp;</p>' * 2)
                html_output.append(
                    f'<h2 style="padding-top:150px; margin-top:-150px;">'
                    f'<span style="color:#0066CC;"><span style="font-size:20px;"><strong>{content_html}</strong></span></span>'
                    f'</h2>'
                )
                html_output.append('<p>&nbsp;</p>')
                last_was_blank = True
                continue

            # H3
            if 'heading 3' in style:
                if not last_was_blank:
                    html_output.append('<p>&nbsp;</p>')
                html_output.append(
                    f'<h3 style="padding-top:150px; margin-top:-150px;">'
                    f'<span style="color:#000000;"><span style="font-size:18px;"><strong>{content_html}</strong></span></span>'
                    f'</h3>'
                )
                html_output.append('<p>&nbsp;</p>')
                last_was_blank = True
                continue

            # 一般內文
            html_output.append(f'<p>{content_html}</p>')
            if not (('heading 2' in next_style) or ('heading 3' in next_style) or (next_tag == "tbl")):
                html_output.append('<p>&nbsp;</p>')
                last_was_blank = True
            else:
                last_was_blank = False

        elif tag == "tbl":
            if in_list: # <--- 插入這兩行
                html_output.append("</ol>")
                in_list = False
            for tbl in doc.tables:
                if tbl._element == element:
                    if not last_was_blank:
                        html_output.append('<p>&nbsp;</p>')
                    html_output.append(table_to_html(tbl))
                    html_output.append('<p>&nbsp;</p>')
                    last_was_blank = True
                    break
    # 確保文件結束時清單已關閉
    if in_list:
        html_output.append("</ol>")
    
    footer_html = CATEGORY_TO_FOOTER_HTML.get(category_choice, "")
    html_output.append('<p>&nbsp;</p><p>&nbsp;</p>' + footer_html + '<p>&nbsp;</p>')
    html_output = apply_auto_toc_and_smooth(html_output)
    result = "\n".join(html_output)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".txt")
    with open(tmp.name, "w", encoding="utf-8") as f:
        f.write(result)

    return tmp.name, result, h1_text or "（文件中無 H1 標題）"

# ==========================================
# 2. 圖片 Resize + 壓縮工具
# ==========================================

def process_image_action(img_file, width, height, quality):
    if img_file is None:
        gr.Info("⚠️ 請先上傳圖片")
        return None, None

    img = Image.open(img_file.name)
    img = img.convert("RGB")

    w = int(width) if width else 810
    h = int(height) if height else 540
    img_resized = img.resize((w, h), Image.LANCZOS)

    q = int(quality) if quality else 70

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".jpg")
    img_resized.save(tmp.name, format="JPEG", quality=q, optimize=True)

    size_kb = os.path.getsize(tmp.name) / 1024
    info = f"✅ 已輸出 {w}×{h}，品質 {q}%｜約 {size_kb:.1f} KB"
    return img_resized, info

def clear_image_action():
    gr.Info("🧹 已重置圖片區")
    return None, 810, 540, 70, None, ""

# ==========================================
# 3. Gradio Action 包裝
# ==========================================

def convert_action(input_file, category_choice):
    if input_file is None:
        gr.Info("⚠️ 請先上傳 Word 檔案")
        return None, None, None, None
    path, html, h1_text = docx_to_html_with_links(input_file, category_choice)
    gr.Info("✅ 轉換成功！")
    return path, html, h1_text, html

def clear_action():
    gr.Info("🧹 已重置")
    return None, None, None, None, None, None

# ==========================================
# 4. UI + CSS
# ==========================================

theme = gr.themes.Soft(
    primary_hue="blue",
    neutral_hue="slate",
    font=[gr.themes.GoogleFont("Noto Sans TC"), "sans-serif"]
)

css = """
/* 整體深色背景 */
body, .gradio-container {
    background: radial-gradient(
        circle at top left,
        #1f2937 0,
        #020617 40%,
        #000 100%
    ) !important;
    color: #e5e7eb !important;
}

/* Panel / Box */
.gr-panel, .gr-box, .gr-group, .gr-form,
.gr-column > .container, .gr-row > .container {
    background: rgba(15, 23, 42, 0.92) !important;
    border-radius: 18px !important;
    border: 1px solid rgba(148, 163, 184, 0.35) !important;
    box-shadow: 0 18px 40px rgba(15, 23, 42, 0.65) !important;
    overflow: visible !important;
}

/* 區塊標題 */
.gr-box > .gr-markdown h3,
.gr-box > .gr-markdown h2 {
    color: #e5e7eb !important;
}

/* Label 玻璃效果 */
label[data-testid="block-label"],
.gr-file > label,
.label-wrap > label,
.form-label > label,
.gr-form > label {
    background: rgba(30, 41, 59, 0.38) !important;
    padding: 6px 14px !important;
    border-radius: 12px !important;
    color: #e5e7eb !important;
    font-weight: 600 !important;
    font-size: 14px !important;
    letter-spacing: 0.02em;
    border: 1px solid rgba(148, 163, 184, 0.25) !important;
    backdrop-filter: blur(6px) !important;
    -webkit-backdrop-filter: blur(6px) !important;
    box-shadow: 0 4px 14px rgba(59,130,246,0.18) !important;
}
.gr-markdown h1 label,
.gr-markdown h2 label,
.gr-markdown h3 label {
    background: none !important;
    box-shadow: none !important;
    border: none !important;
}
span[data-testid="block-info"] {
    background: rgba(30, 41, 59, 0.42) !important;
    color: #e5e7eb !important;
    padding: 6px 14px !important;
    border-radius: 12px !important;
    font-weight: 600 !important;
    font-size: 14px !important;
    letter-spacing: .02em;
    border: 1px solid rgba(148, 163, 184, .25) !important;
    backdrop-filter: blur(6px) !important;
    box-shadow: 0 4px 14px rgba(0,0,0,0.20) !important;
}
.wrap.svelte-1hfxprf.container {
    background: rgba(15, 23, 42, 0.25) !important;
    border: 1px solid rgba(71, 85, 105, .35) !important;
    border-radius: 12px !important;
    padding: 6px 6px !important;
    backdrop-filter: blur(4px) !important;
}

/* 移除子物件捲軸 */
.gradio-container * {
    scrollbar-width: none !important;
}
.gradio-container *::-webkit-scrollbar {
    width: 0 !important;
    height: 0 !important;
}

/* Preview 區 */
#preview-box, #code-box .cm-scroller, #img-preview-box {
    height: 600px !important;
    max-height: 600px !important;
    overflow-y: auto !important;
    border-radius: 14px !important;
    border: 1px solid rgba(55, 65, 81, 0.95) !important;
    background: radial-gradient(
        circle at top left,
        #0f172a 0,
        #020617 55%,
        #020617 100%
    ) !important;
    padding: 16px !important;
    color: #e5e7eb !important;
    font-size: 14px;
}
#img-preview-box img {
    max-width: 100%;
    height: auto;
    border-radius: 12px;
}

/* 表單欄位 */
.gradio-container .gr-input,
.gradio-container .gr-select,
.gradio-container .gr-file {
    background-color: rgba(15, 23, 42, 0.95) !important;
    border-radius: 12px !important;
    border: 1px solid rgba(71, 85, 105, 0.9) !important;
    color: #e5e7eb !important;
}

/* 按鈕樣式 */
#convert-btn, #clear-btn, #img-convert-btn, #img-clear-btn {
    position: relative;
    overflow: hidden;
    border-radius: 999px !important;
    padding: 0.6rem 1.4rem !important;
    font-weight: 700 !important;
    letter-spacing: 0.04em;
    transition: all 0.22s ease-out;
}
#convert-btn, #img-convert-btn {
    border: 1px solid rgba(96, 165, 250, 0.7) !important;
    color: #e5e7eb !important;
    background: radial-gradient(
        circle at 0% 0%,
        rgba(56, 189, 248, 0.85) 0,
        rgba(37, 99, 235, 0.95) 40%,
        rgba(15, 23, 42, 1) 100%
    ) !important;
    box-shadow:
        0 0 0 1px rgba(15, 23, 42, 0.9),
        0 12px 30px rgba(37, 99, 235, 0.55);
}
#convert-btn:hover, #img-convert-btn:hover {
    box-shadow:
        0 0 0 1px rgba(191, 219, 254, 0.9),
        0 18px 45px rgba(56, 189, 248, 0.75);
    transform: translateY(-1px) scale(1.02);
}
#clear-btn, #img-clear-btn {
    border: 1px solid rgba(148, 163, 184, 0.7) !important;
    color: #e5e7eb !important;
    background: linear-gradient(
        135deg,
        rgba(31, 41, 55, 0.95),
        rgba(15, 23, 42, 1)
    ) !important;
    box-shadow:
        0 0 0 1px rgba(15, 23, 42, 1),
        0 10px 26px rgba(15, 23, 42, 0.9);
}
#clear-btn:hover, #img-clear-btn:hover {
    border-color: rgba(209, 213, 219, 0.95) !important;
    transform: translateY(-1px);
}
#convert-btn::before,
#clear-btn::before,
#img-convert-btn::before,
#img-clear-btn::before {
    content: "";
    position: absolute;
    top: 0;
    left: -120%;
    width: 100%;
    height: 100%;
    background: linear-gradient(
        120deg,
        transparent 0%,
        rgba(255, 255, 255, 0.25) 40%,
        rgba(255, 255, 255, 0.75) 50%,
        rgba(255, 255, 255, 0.25) 60%,
        transparent 100%
    );
    opacity: 0;
}
#convert-btn:hover::before,
#clear-btn:hover::before,
#img-convert-btn:hover::before,
#img-clear-btn:hover::before {
    opacity: 1;
    animation: shimmer-slide 0.9s ease-out forwards;
}
@keyframes shimmer-slide {
    0% { transform: translateX(0); left: -120%; }
    100% { transform: translateX(120%); left: 120%; }
}
#left-panel { padding: 18px 20px !important; }
#right-panel { padding: 18px 22px !important; }

#left-panel-img { padding: 18px 20px !important; }
#right-panel-img { padding: 18px 22px !important; }

/* 移除圖片工具分享按鈕 */
#img-preview-box .svelte-1ipelgc:nth-of-type(3),
#img-preview-box button[aria-label="Share"] {
    display: none !important;
}
#img-preview-box .svelte-1ipelgc {
    opacity: 1 !important;
}
#img-preview-box img {
    image-rendering: auto;
}
"""

# ==========================================
# 5. Gradio 介面
# ==========================================

with gr.Blocks(theme=theme, css=css, title="i-Buzz Editor 文章轉換器") as demo:

    with gr.Row():
        with gr.Column():
            gr.Markdown("## 🐝 i-Buzz Editor 文章轉換器")

            with gr.Accordion("📘 操作說明（點擊展開）", open=False):

                gr.HTML("""
                <div style="line-height:1.7; font-size:16px;">

                  <h2 style="margin-bottom:10px;">🛠️ 這個工具能幫你做什麼？</h2>

                  <ul style="margin-left:18px;">
                    <li><strong>自動把 Word 原稿轉成官網可用的 HTML</strong>（標題階層、段落、空行全部重整）</li>
                    <li><strong>自動把獨立一行的 URL 轉成嵌入卡片</strong>（IG / Threads / FB / YouTube）</li>
                    <li><strong>支援智慧判斷</strong>（例如 IG Reel、圖片帖、FB 影片等會自動調整嵌入高度）</li>
                    <li><strong>Word 表格 → 完整 HTML 表格</strong>（邊框、粗體、自動排版）</li>
                    <li><strong>圖片壓縮＋調整尺寸</strong>（第二個分頁可一次處理）</li>
                    <li><strong>自動加上 Footer CTA</strong>（依分類套用 i-Buzz / AsiaKOL / FansFeed / Threads…）</li>
                  </ul>

                  <br>

                  <h2 style="margin-bottom:10px;">📌 開始前一定要確認的 3 件事</h2>
                  <ul style="margin-left:18px; list-style-type: square;">
                    <li><strong>標題階層要正確：</strong>H1=主標、H2=大標、H3=小標（H1 會自動抽出，不顯示在文章內）</li>
                    <li><strong>網址要獨立成一行：</strong>整行只有 URL 才會轉成卡片</li>
                    <li><strong>空行不用手動調整：</strong>系統會自動調整漂亮排版</li>
                  </ul>

                  <br>

                  <h2 style="margin-bottom:10px;">🚀 轉檔步驟</h2>
                  <ol style="margin-left:18px;">
                    <li>上傳 <code>.docx</code> 原稿</li>
                    <li>選擇文章分類（會自動套用對應 Footer）</li>
                    <li>按下「開始轉換」</li>
                    <li>右側可預覽、可複製，也能直接下載 HTML 檔</li>
                  </ol>

                  <br>

                  <h2 style="margin-bottom:6px;">📄 範例原稿下載（Demo）</h2>
                  <p>以下是「建議格式」的示範檔，你可以下載照著排，轉檔最穩定：</p>

                  <a href="https://docs.google.com/document/d/1lUKgxM--8VeTYHpvX7hdlb19toZ2VYh8/export?format=docx"
                     download
                     style="
                        display:inline-block;
                        padding:10px 20px;
                        margin-top:8px;
                        border-radius:12px;
                        background:linear-gradient(135deg, #3b82f6 0%, #2563eb 100%);
                        color:white;
                        font-weight:600;
                        text-decoration:none;
                        letter-spacing:0.03em;
                        box-shadow:0 4px 14px rgba(37, 99, 235, 0.35);
                     ">
                     📎 點我下載 converter_demo.docx（範例原稿）
                  </a>

                  <br><br>

                </div>
                """)

    with gr.Tabs():
        ...



        # 文章轉檔
        with gr.TabItem("📝 文章轉檔"):
            with gr.Row(equal_height=False):
                with gr.Column(scale=1, elem_id="left-panel"):
                    gr.Markdown("### 🔧 設定與動作")

                    file_input = gr.File(
                        label="📂 上傳 Word 檔（.docx）",
                        file_types=[".docx"],
                        file_count="single"
                    )

                    category_choice = gr.Dropdown(
                        choices=[
                            "🔵 數據分析解方",
                            "🔷 產業口碑數據",
                            "🟦 消費者洞察",
                            "🩷 網紅行銷策略",
                            "🟡 社群粉絲團健檢",
                            "🟣 Threads 行銷服務"
                        ],
                        label="#️⃣文章分類（決定 Footer）",
                        value="🔵 數據分析解方",
                        interactive=True
                    )

                    with gr.Row():
                        convert_btn = gr.Button("✨ 開始轉換", variant="primary", elem_id="convert-btn")
                        clear_btn = gr.Button("🧹 重置", elem_id="clear-btn")

                    download_output = gr.File(visible=False)

                with gr.Column(scale=3, elem_id="right-panel"):
                    gr.Markdown("### 📄 轉換結果")

                    h1_output = gr.Textbox(
                        label="🔖主標題（H1）",
                        interactive=False,
                        show_copy_button=True
                    )

                    with gr.Tabs():
                        with gr.TabItem("🌐 HTML 預覽"):
                            html_preview = gr.HTML(label="HTML Preview", elem_id="preview-box")
                        with gr.TabItem("💻 HTML 原始碼"):
                            code_output = gr.Code(
                                language="html",
                                label="HTML Code",
                                interactive=False,
                                elem_id="code-box"
                            )

            convert_btn.click(
                fn=convert_action,
                inputs=[file_input, category_choice],
                outputs=[download_output, html_preview, h1_output, code_output]
            )

            clear_btn.click(
                fn=clear_action,
                inputs=None,
                outputs=[file_input, download_output, category_choice, html_preview, h1_output, code_output]
            )

        # 圖片工具
        with gr.TabItem("🖼️ 圖片壓縮調整大小"):
            with gr.Row(equal_height=False):
                with gr.Column(scale=1, elem_id="left-panel-img"):
                    gr.Markdown("### 🔧 圖片設定與動作")

                    img_input = gr.File(
                        label="📎 上傳圖片（jpg / png / webp）",
                        file_types=[".jpg", ".jpeg", ".png", ".webp"],
                        file_count="single"
                    )

                    width_in = gr.Number(label="寬度（px）", value=810, precision=0)
                    height_in = gr.Number(label="高度（px）", value=540, precision=0)

                    quality_in = gr.Slider(
                        minimum=30, maximum=95, value=70, step=1,
                        label="壓縮品質（%）"
                    )

                    with gr.Row():
                        img_convert_btn = gr.Button("✨ 開始處理", variant="primary", elem_id="img-convert-btn")
                        img_clear_btn = gr.Button("🧹 重置", elem_id="img-clear-btn")

                with gr.Column(scale=3, elem_id="right-panel-img"):
                    gr.Markdown("### 👀 圖片預覽")
                    img_preview = gr.Image(
                        label="Preview",
                        elem_id="img-preview-box",
                        format="jpeg"
                    )

                    img_info = gr.Markdown("")

            img_convert_btn.click(
                fn=process_image_action,
                inputs=[img_input, width_in, height_in, quality_in],
                outputs=[img_preview, img_info]
            )

            img_clear_btn.click(
                fn=clear_image_action,
                inputs=None,
                outputs=[img_input, width_in, height_in, quality_in, img_preview, img_info]
            )

if __name__ == "__main__":
    os.environ["GRADIO_ANALYTICS_ENABLED"] = "False"
    demo.launch(show_error=True, ssr_mode=False)
